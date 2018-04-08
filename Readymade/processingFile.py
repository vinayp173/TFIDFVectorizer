from base64 import _85encode

import docx
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams, LTTextBox, LTTextLine
from nltk.corpus import stopwords
import nltk
from nltk.tag import *


# from lib2to3.fixes.fix_input import context

# Extracting Words from .txt document
def Txt_generation_TF(f, qaStatus=False):
    num_words = 0
    allSentances = []
    NativeallWords = []
    for line in f:
        line = u'' + line
        dummy = line.replace('\n', '').replace('/', '')
        allSentances.extend(dummy.split('.'))
        words = line.split()
        NativeallWords.extend(words)
        num_words += len(words)
    allSentances = list(filter(None, allSentances))
    # print(allSentances)
    if qaStatus:
        print('from text ')
        return allSentances

    return calc_TF(NativeallWords, num_words)


# Extracting Words from .pdf document
def Pdf_generation_TF(f, qaStatus=False):
    allSentances = []
    num_words = 0
    parser = PDFParser(f)
    doc = PDFDocument()
    parser.set_document(doc)
    doc.set_parser(parser)
    doc.initialize('')
    rsrcmgr = PDFResourceManager()
    laparams = LAParams()
    device = PDFPageAggregator(rsrcmgr, laparams=laparams)
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    # Process each page contained in the document.
    NativeallWords = []
    for page in doc.get_pages():
        interpreter.process_page(page)
        layout = device.get_result()
        for lt_obj in layout:
            if isinstance(lt_obj, LTTextBox) or isinstance(lt_obj, LTTextLine):
                num_words += len(lt_obj.get_text().split())
                lt_obj.get_text().encode("utf8")
                rawSentance = str(lt_obj)
                # allSentances=sentanceGernaration(rawSentance)
                endPoint = rawSentance.rfind("\\n") - 1
                rawSentance = rawSentance[55:endPoint].replace('\\n', '').replace('\\s', '')
                allSentances.extend(rawSentance.split('.'))
                NativeallWords.extend(lt_obj.get_text().split())

    # print(NativeallWords)
    # print('Sentances from pdf',allSentances)
    if qaStatus:
        print('from PDF ')
        return allSentances
    return calc_TF(NativeallWords, num_words)


# Extracting Words from .docx document
def Docx_generation_TF(path, qaStatus=False):
    doc = docx.Document(path)
    NativeallWords = []
    allSentances = []

    #     fname=doc.name
    # print(len(doc.paragraphs))
    paraNo = 0
    num_words = 0
    no_of_paragraphs = len(doc.paragraphs)
    for para in doc.paragraphs:
        num_words += len(para.text.split())
        allSentances.extend(para.text.split('.'))
        NativeallWords.extend(para.text.split())
    # for line in allSentances:
    #     print(line,"*")
    if qaStatus:
        print('from DOCS ')
        return allSentances
    return calc_TF(NativeallWords, num_words)


# Calculating Term Frequincies
def calc_TF(NativeallWords, num_words):
    context = {}
    NativeallWords = removeStopWords(NativeallWords)
    # print(len(NativeallWords))
    Lower_allWords = [w.lower() for w in NativeallWords]
    Lower_allWords = [w.replace(",", " ") for w in Lower_allWords]
    Lower_allWords = [w.replace(".", " ") for w in Lower_allWords]
    Lower_allWords = [w.split("'")[0] for w in Lower_allWords]

    countWord = 0
    tf = 0
    for w in Lower_allWords:
        if not w.isdigit() and (w.isalpha() or w.isalnum()) and len(w) < 50:
            countWord = Lower_allWords.count(w)
            tf = (countWord / num_words)
            context[w] = tf
    print('TF calculation Done')

    return context


def removeStopWords(NativeallWords):
    nltk.download('stopwords')
    nltk.download('punkt')
    # stop = set(stopwords.words('english'))
    # list = [w for w in NativeallWords if w not in stop]
    # nltk.download('averaged_perceptron_tagger')
    # keylist = pos_tag(NativeallWords)
    # list = [word for word, pos in keylist if
    #         pos != 'VBZ' and pos != 'DT' and pos != 'IN' and pos != 'PRP' and pos != 'CC']
    stemmer = nltk.SnowballStemmer("english")
    list = [stemmer.stem(w).lower() for w in NativeallWords]
    return list

# temporary function for testing only
# def dummyNLTK(NativeallWords):
#     NativeallWords = removeStopWords(NativeallWords)
#     Lower_allWords = [w.lower() for w in NativeallWords]
#     Lower_allWords = [w.replace(",", " ") for w in Lower_allWords]
#     Lower_allWords = [w.replace(".", " ") for w in Lower_allWords]
#     Lower_allWords = [w.split("'")[0] for w in Lower_allWords]
#     print(len(Lower_allWords),"<-----Count With Nltk")
